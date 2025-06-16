import io
import re
from pdfminer.converter import TextConverter
from pdfminer.pdfinterp import PDFPageInterpreter
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from docx import Document
import spacy
from spacy.matcher import Matcher
import os # Added for __main__ example cleanup

# nlp object will be loaded in the parse_resume function when first needed
nlp = None

def extract_text_from_pdf(pdf_path):
    """Extracts text from a PDF file."""
    resource_manager = PDFResourceManager()
    fake_file_handle = io.StringIO()
    converter = TextConverter(resource_manager, fake_file_handle, laparams=LAParams())
    page_interpreter = PDFPageInterpreter(resource_manager, converter)

    try:
        with open(pdf_path, 'rb') as fh:
            for page in PDFPage.get_pages(fh,
                                          caching=True,
                                          check_extractable=True):
                page_interpreter.process_page(page)
            text = fake_file_handle.getvalue()
    except Exception as e:
        print(f"Error reading PDF file {pdf_path}: {e}")
        return None
    finally:
        converter.close()
        fake_file_handle.close()

    return text

def extract_text_from_docx(docx_path):
    """Extracts text from a DOCX file."""
    try:
        doc = Document(docx_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except Exception as e:
        print(f"Error reading DOCX file {docx_path}: {e}")
        return None

def extract_text_from_file(file_path):
    """Extracts text from PDF or DOCX file based on extension."""
    if not file_path: # Handle cases where file_path might be None
        print("Error: No file path provided for text extraction.")
        return None
    if file_path.endswith(".pdf"):
        return extract_text_from_pdf(file_path)
    elif file_path.endswith(".docx"):
        return extract_text_from_docx(file_path)
    else:
        print(f"Unsupported file format: {file_path}. Only PDF and DOCX are supported.")
        return None

# --- spaCy based extraction functions ---

def extract_name(text_doc): # Expects a spaCy Doc object
    """Extracts name using spaCy's NER (PERSON entities)."""
    if text_doc is None:
        return "Name Not Found (No document processed)"
    for ent in text_doc.ents:
        if ent.label_ == "PERSON":
            # Could add more logic here to find the most likely candidate if multiple PERSON entities exist
            return ent.text.strip()
    return "Name Not Found"

def extract_email(text): # Expects raw text
    """Extracts email using a regex pattern."""
    if text is None:
        return "Email Not Found (No text provided)"
    # Improved regex for common email patterns
    email_regex = r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"
    matches = re.findall(email_regex, text)
    return matches[0] if matches else "Email Not Found"

def extract_phone(text): # Expects raw text
    """Extracts phone number using a more comprehensive regex pattern."""
    if text is None:
        return "Phone Not Found (No text provided)"
    # Regex to find various phone number formats (North American focus, can be adapted)
    phone_regex = r"(\+?\d{1,3}[-\.\s]?)?(\(\d{3}\)|\d{3})[-\.\s]?\d{3}[-\.\s]?\d{4}"
    matches = re.findall(phone_regex, text)
    # Post-process matches to return a clean number if found
    if matches:
        # The regex might return tuples if grouping is used for optional parts like country code
        # We want to join them and return the first full match
        for match in matches:
            if isinstance(match, tuple):
                full_match_str = "".join(m for m in match if m) # Join non-empty parts of the tuple
                # Further clean up common separators from the found string
                return re.sub(r"[^0-9]", "", full_match_str)
            elif isinstance(match, str) and match.strip(): # If it's a direct string match
                 return re.sub(r"[^0-9]", "", match) # Clean it
    return "Phone Not Found"

def extract_education(text_doc): # Expects a spaCy Doc object
    """Extracts education-related text using keyword matching and sentence analysis (rudimentary)."""
    if text_doc is None:
        return "Education Not Found (No document processed)"

    education_keywords = [
        'education', 'university', 'college', 'institute', 'bachelor', 'master', 'phd',
        'degree', 'qualification', 'academic', 'school', 'cgpa', 'gpa'
    ]
    education_section = []

    # Simple approach: look for sentences containing education keywords
    # More advanced: identify sections titled "Education" or similar
    for sent in text_doc.sents:
        sent_text = sent.text.lower()
        if any(keyword in sent_text for keyword in education_keywords):
            education_section.append(sent.text.strip())

    return "\n".join(education_section) if education_section else "Education information not clearly found"

def extract_skills(text_doc, skills_list=None): # Expects a spaCy Doc object
    """Extracts skills using spaCy's Matcher."""
    if text_doc is None:
        return ["Skills Not Found (No document processed)"]

    # Initialize Matcher with the shared vocab
    matcher = Matcher(nlp.vocab)

    # Default skills list if none provided (can be expanded significantly)
    if skills_list is None:
        skills_list = [
            "python", "java", "c++", "javascript", "react", "angular", "vue", "node.js",
            "sql", "nosql", "mongodb", "postgresql", "mysql",
            "aws", "azure", "gcp", "docker", "kubernetes",
            "machine learning", "deep learning", "tensorflow", "pytorch", "keras", "scikit-learn",
            "data analysis", "data science", "pandas", "numpy", "matplotlib", "seaborn",
            "natural language processing", "nlp", "spacy", "nltk",
            "agile", "scrum", "jira", "git", "communication", "problem solving", "teamwork"
        ]

    # Create patterns for the Matcher
    # Simple phrase matching for now. Can be extended with more complex patterns.
    for skill in skills_list:
        # Handle multi-word skills by creating a list of token patterns
        pattern = []
        for token in skill.split():
            pattern.append({"LOWER": token}) # Match lowercased token
        matcher.add(skill.upper(), [pattern]) # Use skill name as pattern ID

    matches = matcher(text_doc)

    found_skills = set() # Use a set to store unique skills
    for match_id, start, end in matches:
        skill_name = nlp.vocab.strings[match_id] # Get the skill name (ID)
        found_skills.add(skill_name.lower()) # Add the matched skill (original from skills_list)

    return list(found_skills) if found_skills else ["No specific skills found from the predefined list"]


def parse_resume(file_path):
    """Main function to parse resume and extract information."""
    # Load spaCy model if not already loaded (critical for other functions)
    global nlp
    if nlp is None:
        try:
            nlp = spacy.load('en_core_web_sm')
            print("spaCy model 'en_core_web_sm' loaded successfully for parsing.")
        except OSError:
            print("CRITICAL: Failed to load spaCy model 'en_core_web_sm'. Make sure it's installed.")
            print("Attempting to download and load en_core_web_sm...")
            try:
                spacy.cli.download('en_core_web_sm')
                nlp = spacy.load('en_core_web_sm')
                print("Successfully downloaded and loaded 'en_core_web_sm'.")
            except Exception as e:
                print(f"Still failed to load spaCy model after download attempt: {e}")
                return {"error": "spaCy model 'en_core_web_sm' not available or failed to download."}
        except Exception as e: # Catch any other spacy.load exception
            print(f"An unexpected error occurred while loading spaCy model: {e}")
            return {"error": "Unexpected error loading spaCy model."}


    text = extract_text_from_file(file_path)
    if not text:
        return {"error": f"Could not extract text from {file_path}"}

    # Process the text with spaCy to get a Doc object
    # This should happen after nlp is confirmed to be loaded
    if nlp is None : # Double check, should not happen if above logic is correct
        return {"error": "spaCy nlp object not initialized before processing text."}

    doc = nlp(text)

    name = extract_name(doc)
    email = extract_email(text) # Email regex works on raw text
    phone = extract_phone(text) # Phone regex works on raw text
    education = extract_education(doc)
    skills = extract_skills(doc) # Pass the Doc object

    return {
        "text": text,
        "name": name,
        "email": email,
        "phone": phone,
        "education": education,
        "skills": skills,
        "no_of_pages": "N/A (pdfminer specific)" # Placeholder, as pyresparser provided this
    }

if __name__ == '__main__':
    # Example Usage (for testing this module directly)
    print("Resume Parser Module - Running Test")
    # Create a dummy PDF for testing if spaCy and pdfminer are working
    # This requires reportlab or similar, which is not a direct dependency yet.
    # For now, this test will rely on manual placement of a test file.

    # Create a dummy docx file for testing
    try:
        from docx import Document as TestDoc # Renamed to avoid conflict with outer scope Document
        test_docx_content = TestDoc()
        test_docx_content.add_paragraph("Dr. Jane Doe jane.doe@example.com (123) 456-7890")
        test_docx_content.add_paragraph("Education: BSc in Computer Science from Tech University. Skills: Python, Java.")
        test_docx_content.save("dummy_test_resume.docx")
        print("Created dummy_test_resume.docx")

        data_docx = parse_resume("dummy_test_resume.docx")
        if data_docx and "error" not in data_docx:
            print("\n--- Parsed DOCX Data (dummy_test_resume.docx) ---")
            for key, value in data_docx.items():
                if key == "text":
                    print(f"{key.capitalize()}: {value[:150]}...")
                else:
                    print(f"{key.capitalize()}: {value}")
        elif data_docx:
            print(f"Error parsing dummy_test_resume.docx: {data_docx.get('error')}")
        else:
            print("Unknown error parsing dummy_test_resume.docx")

        # Clean up dummy file
        # Ensure os is imported if this file is run standalone. It was added to the imports.
        if os.path.exists("dummy_test_resume.docx"):
             os.remove("dummy_test_resume.docx")
             print("Removed dummy_test_resume.docx")

    except ImportError:
        print("Skipping __main__ test for docx: python-docx not found in this subtask environment for test file creation.")
    except Exception as e:
        print(f"Error in __main__ test block: {e}")
